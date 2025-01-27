from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from datetime import datetime

def coletar_dados_usuario():
    """Coleta os dados básicos do usuário."""
    print('*' * 6 + ' Sistema de cadastro de Pronto Socorro ' + '*' * 6)
    data_hora_ocorrencia = input("Digite a data e hora da ocorrência (dd/mm/aaaa hh:mm): ")
    nome_usuario = input("Digite o nome do usuário(a): ")
    cpf_usuario = input("Digite o CPF do usuário(a): ")
    endereco_usuario = input("Digite o endereço do usuário(a): ")
    local_ocorrencia = input("Digite o local da ocorrência: ")
    descricao_ocorrencia = input("Descreva a ocorrência: ")

    return {
        "data_hora_ocorrencia": data_hora_ocorrencia,
        "nome_usuario": nome_usuario,
        "cpf_usuario": cpf_usuario,
        "endereco_usuario": endereco_usuario,
        "local_ocorrencia": local_ocorrencia,
        "descricao_ocorrencia": descricao_ocorrencia,
    }


def coletar_dados_socorro():
    """Coleta os dados adicionais caso o socorro seja aceito."""
    atendentes = input("Informe os atendentes (nome, RE, separados por vírgula): ")
    linha_coletivo = input("Informe a linha do coletivo: ")
    motorista = input("Informe o nome do motorista: ")
    chapa_motorista = input("Digite a chapa do motorista: ")
    cobrador = input("Informe o nome do cobrador: ")
    chapa_cobrador = input("Informe a chapa do cobrador: ")
    destino_usuario = input("Informe o destino do usuário após a ocorrência: ")
    data_hora_encerramento = input("Informe o horário e a data do encerramento da ocorrência: ")
    retorno_funcionarios = input("Informe a data e hora do retorno dos funcionários: ")

    return {
        "atendentes": atendentes,
        "linha_coletivo": linha_coletivo,
        "motorista": motorista,
        "chapa_motorista": chapa_motorista,
        "cobrador": cobrador,
        "chapa_cobrador": chapa_cobrador,
        "destino_usuario": destino_usuario,
        "data_hora_encerramento": data_hora_encerramento,
        "retorno_funcionarios": retorno_funcionarios,
    }


def salvar_relatorio_docx(relatorio):
    """Salva o relatório em um arquivo .docx com timestamp no nome."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nome_arquivo = f'relatorio_ocorrencia_{timestamp}.docx'
    doc = Document()
    doc.add_heading('Relatório de Ocorrência', level=1)
    doc.add_paragraph(relatorio)
    doc.save(nome_arquivo)
    print(f"Relatório salvo com sucesso em '{nome_arquivo}'!")
    return nome_arquivo


def enviar_email_gmail(destinatario, assunto, conteudo, anexo):
    """Envia o e-mail usando o servidor SMTP do Gmail."""
    # Configuração do Gmail
    remetente = "franco.jose.santos@gmail.com"  # Substitua pelo seu e-mail
    senha = "smxyrzchqenvuorv"  # Substitua pela sua senha de app do Gmail

    try:
        # Configurando o servidor SMTP
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()  # Inicia o TLS para segurança
        server.login(remetente, senha)

        # Criando a mensagem
        mensagem = MIMEMultipart()
        mensagem['From'] = remetente
        mensagem['To'] = destinatario
        mensagem['Subject'] = assunto

        # Corpo do e-mail
        mensagem.attach(MIMEText(conteudo, 'plain'))

        # Anexo
        with open(anexo, 'rb') as arquivo:
            parte_anexo = MIMEBase('application', 'octet-stream')
            parte_anexo.set_payload(arquivo.read())
            encoders.encode_base64(parte_anexo)
            parte_anexo.add_header(
                'Content-Disposition',
                f'attachment; filename={anexo}'
            )
            mensagem.attach(parte_anexo)

        # Enviando o e-mail
        server.send_message(mensagem)
        print("E-mail enviado com sucesso!")
    except Exception as e:
        print(f"Erro ao enviar o e-mail: {e}")
    finally:
        server.quit()


def cadastrar_ocorrencia():
    dados_usuario = coletar_dados_usuario()

    aceitar_socorro = input("Usuário(a) aceita socorro? (sim/não): ").lower()
    if aceitar_socorro in ['não', 'nao']:
        print("Socorro negado pelo usuário. Encerrando o registro de ocorrência.")
        return

    dados_socorro = coletar_dados_socorro()

    relatorio = f"""
Ocorrência registrada:

Às {dados_usuario['data_hora_ocorrencia']}, usuária {dados_usuario['nome_usuario']}, CPF: {dados_usuario['cpf_usuario']}, residente na {dados_usuario['endereco_usuario']}, encontrava-se pela {dados_usuario['local_ocorrencia']} onde veio acometer-se de mal súbito ({dados_usuario['descricao_ocorrencia']}). Realizados primeiros atendimentos pelos O.T.U'S {dados_socorro['atendentes']}, que após verificação do ocorrido, conduziu a usuária com apoio da cadeira de rodas, dentro do coletivo da linha {dados_socorro['linha_coletivo']}, motorista responsável: {dados_socorro['motorista']}, Chapa: {dados_socorro['chapa_motorista']} e cobrador {dados_socorro['cobrador']}, Chapa: {dados_socorro['chapa_cobrador']}. Às {dados_socorro['data_hora_encerramento']} conduzida ao {dados_socorro['destino_usuario']}, onde ficou sob cuidados médicos. Ocorrência encerrada às {dados_socorro['data_hora_encerramento']}. Retorno dos funcionários às {dados_socorro['retorno_funcionarios']}.
"""
    nome_arquivo = salvar_relatorio_docx(relatorio)

    # Enviar o relatório por e-mail
    destinatario = input("Digite o e-mail do destinatário: ")
    assunto = "Relatório de Ocorrência"
    conteudo = "Segue em anexo o relatório de ocorrência."
    enviar_email_gmail(destinatario, assunto, conteudo, nome_arquivo)


# Chamando a função para cadastrar uma ocorrência
cadastrar_ocorrencia()
